from sqlalchemy.orm import Session
from app.models.erp_models import (
    Product, LeadProduct, InvoiceItem, Payment,
    ContractTemplate, Contract, Activity, Attachment
)
from app.models.lead_models import Lead
from app.models.invoice_model import Invoice
from app.schemas.erp_schemas import ProductCreate, ProductUpdate, LeadProductCreate, PaymentCreate
from datetime import datetime, timedelta
from typing import List, Optional
import os, re

# ─────────────── PRODUCT SERVICE ───────────────
class ERPService:

    @staticmethod
    def get_products(db: Session):
        return db.query(Product).order_by(Product.created_at.desc()).all()

    @staticmethod
    def create_product(db: Session, product_data: ProductCreate):
        data = product_data.model_dump()
        db_product = Product(**data)
        db.add(db_product)
        db.commit()
        db.refresh(db_product)
        return db_product

    @staticmethod
    def update_product(db: Session, product_id: int, product_data: ProductUpdate):
        db_product = db.query(Product).filter(Product.id == product_id).first()
        if not db_product:
            return None
        for key, value in product_data.model_dump(exclude_unset=True).items():
            setattr(db_product, key, value)
        db.commit()
        db.refresh(db_product)
        return db_product

    # ─────────────── LEAD PRODUCTS ───────────────
    @staticmethod
    def add_product_to_lead(db: Session, lead_product_data: LeadProductCreate):
        data = lead_product_data.model_dump()
        db_lp = LeadProduct(**data)
        db.add(db_lp)
        ERPService._log(db, lead_id=data["lead_id"], action="Product Added",
                        entity_type="LeadProduct", entity_id=0,
                        details=f"Added product: {data.get('product_name', '')}")
        db.commit()
        db.refresh(db_lp)
        return db_lp

    @staticmethod
    def get_lead_products(db: Session, lead_id: int):
        return db.query(LeadProduct).filter(LeadProduct.lead_id == lead_id).all()

    @staticmethod
    def save_lead_products(db: Session, lead_id: int, products: List[dict]):
        """Replace all lead products with new list (upsert behavior)."""
        # Delete existing
        db.query(LeadProduct).filter(LeadProduct.lead_id == lead_id).delete()
        # Insert new
        for p in products:
            qty = int(p.get("quantity", 1) or 1)
            price = float(p.get("unit_price", 0) or 0)
            tax = float(p.get("tax", 0) or 0)
            disc = float(p.get("discount", 0) or 0)
            base = price * qty
            after_disc = base - (base * disc / 100)
            total = after_disc * (1 + tax / 100)

            lp = LeadProduct(
                lead_id=lead_id,
                product_id=p.get("product_id"),
                product_name=p.get("name") or p.get("product_name") or "",
                description=p.get("description", ""),
                quantity=qty,
                unit_price=price,
                billing_cycle=p.get("billing_cycle", "One-time"),
                start_date=p.get("start_date") or None,
                end_date=p.get("end_date") or None,
                tax=tax,
                discount=disc,
                total=round(total, 2)
            )
            db.add(lp)

        ERPService._log(db, lead_id=lead_id, action="Products Updated",
                        entity_type="Lead", entity_id=lead_id,
                        details=f"Products list updated with {len(products)} item(s)")
        db.commit()
        return ERPService.get_lead_products(db, lead_id)

    # ─────────────── INVOICE SERVICE ───────────────
    @staticmethod
    def get_invoices(db: Session, lead_id: Optional[int] = None):
        q = db.query(Invoice)
        if lead_id:
            q = q.filter(Invoice.lead_id == lead_id)
        invoices = q.order_by(Invoice.created_at.desc()).all()
        result = []
        for inv in invoices:
            items = db.query(InvoiceItem).filter(InvoiceItem.invoice_id == inv.id).all()
            lead = db.query(Lead).filter(Lead.id == inv.lead_id).first() if inv.lead_id else None
            inv_dict = {
                "id": inv.id,
                "invoice_number": inv.invoice_number,
                "lead_id": inv.lead_id,
                "lead_name": lead.lead_name if lead else None,
                "amount": float(inv.amount or 0),
                "subtotal": float(inv.subtotal or 0),
                "tax_amount": float(inv.tax_amount or 0),
                "discount_amount": float(inv.discount_amount or 0),
                "total_amount": float(inv.total_amount or 0),
                "grand_total": float(inv.total_amount or 0),
                "status": inv.status,
                "billing_type": inv.billing_type,
                "currency": inv.currency,
                "issue_date": inv.issue_date.isoformat() if inv.issue_date else None,
                "due_date": inv.due_date.isoformat() if inv.due_date else None,
                "created_at": inv.created_at.isoformat() if inv.created_at else None,
                "notes": inv.notes,
                "items": [
                    {
                        "id": it.id,
                        "item_name": it.item_name,
                        "description": it.description,
                        "quantity": it.quantity,
                        "unit_price": float(it.unit_price or 0),
                        "tax_amount": float(it.tax_amount or 0),
                        "discount_amount": float(it.discount_amount or 0),
                        "total": float(it.total or 0)
                    } for it in items
                ]
            }
            result.append(inv_dict)
        return result

    @staticmethod
    def generate_invoice(db: Session, lead_id: int, billing_type: str = "One-time"):
        lead = db.query(Lead).filter(Lead.id == lead_id).first()
        if not lead:
            return None

        lead_products = db.query(LeadProduct).filter(LeadProduct.lead_id == lead_id).all()
        # Filter by billing type if needed (allow all for manual generation)
        if not lead_products:
            return None

        subtotal = sum(float(lp.unit_price or 0) * int(lp.quantity or 1) for lp in lead_products)
        tax = sum(float(lp.unit_price or 0) * int(lp.quantity or 1) * (float(lp.tax or 0) / 100) for lp in lead_products)
        discount = sum(float(lp.unit_price or 0) * int(lp.quantity or 1) * (float(lp.discount or 0) / 100) for lp in lead_products)
        total_amount = subtotal + tax - discount

        ts = datetime.now().strftime("%Y%m%d%H%M%S")
        invoice_number = f"INV-{lead_id}-{ts}"

        db_invoice = Invoice(
            lead_id=lead_id,
            customer_id=lead_id,
            invoice_number=invoice_number,
            invoice_month=datetime.now().strftime("%B %Y"),
            amount=subtotal,
            subtotal=subtotal,
            tax_amount=round(tax, 2),
            discount_amount=round(discount, 2),
            total_amount=round(total_amount, 2),
            status="Generated",
            billing_type=billing_type,
            currency="PKR",
            issue_date=datetime.now(),
            due_date=datetime.now() + timedelta(days=15),
            notes=""
        )
        db.add(db_invoice)
        db.flush()

        for lp in lead_products:
            base = float(lp.unit_price or 0) * int(lp.quantity or 1)
            item_tax = base * (float(lp.tax or 0) / 100)
            item_disc = base * (float(lp.discount or 0) / 100)
            item_total = base + item_tax - item_disc
            item = InvoiceItem(
                invoice_id=db_invoice.id,
                product_id=lp.product_id,
                item_name=lp.product_name or "",
                description=lp.description or "",
                quantity=lp.quantity,
                unit_price=float(lp.unit_price or 0),
                tax_amount=round(item_tax, 2),
                discount_amount=round(item_disc, 2),
                total=round(item_total, 2)
            )
            db.add(item)

        lead.invoice_status = "Generated"
        ERPService._log(db, lead_id=lead_id, action="Invoice Generated",
                        entity_type="Invoice", entity_id=db_invoice.id,
                        details=f"Invoice {invoice_number} generated. Total: PKR {total_amount:,.2f}")
        db.commit()
        db.refresh(db_invoice)
        return db_invoice



    @staticmethod
    def generate_invoice_pdf(db: Session, invoice_id: int) -> Optional[str]:
        """Generate a PDF invoice using reportlab if available, else create a basic text file."""
        inv = db.query(Invoice).filter(Invoice.id == invoice_id).first()
        if not inv:
            return None

        items = db.query(InvoiceItem).filter(InvoiceItem.invoice_id == invoice_id).all()
        lead = db.query(Lead).filter(Lead.id == inv.lead_id).first() if inv.lead_id else None

        os.makedirs("storage/invoices", exist_ok=True)
        pdf_path = f"storage/invoices/invoice_{invoice_id}.pdf"

        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib import colors
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
            from reportlab.lib.styles import ParagraphStyle

            doc = SimpleDocTemplate(pdf_path, pagesize=A4,
                                    rightMargin=0.5*inch, leftMargin=0.5*inch,
                                    topMargin=0.75*inch, bottomMargin=0.5*inch)
            styles = getSampleStyleSheet()
            story = []

            # Header
            title_style = ParagraphStyle("title", parent=styles["Heading1"],
                                          fontSize=22, textColor=colors.HexColor("#1e3a5f"),
                                          spaceAfter=4)
            story.append(Paragraph("INVOICE", title_style))
            story.append(Paragraph(f"<font size='10' color='#64748b'>Invoice Number: <b>{inv.invoice_number}</b></font>", styles["Normal"]))
            story.append(Spacer(1, 12))

            # Lead info
            lead_name = lead.lead_name if lead else "N/A"
            company = lead.company_name if lead else "N/A"
            info_data = [
                ["Bill To:", "Invoice Details:"],
                [lead_name, f"Date: {inv.issue_date.strftime('%d %b %Y') if inv.issue_date else 'N/A'}"],
                [company, f"Due: {inv.due_date.strftime('%d %b %Y') if inv.due_date else 'N/A'}"],
                ["", f"Status: {inv.status}"],
            ]
            info_table = Table(info_data, colWidths=[3*inch, 3*inch])
            info_table.setStyle(TableStyle([
                ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
                ("FONTSIZE", (0,0), (-1,-1), 10),
                ("TEXTCOLOR", (0,0), (-1,0), colors.HexColor("#1e3a5f")),
                ("VALIGN", (0,0), (-1,-1), "TOP"),
                ("BOTTOMPADDING", (0,0), (-1,-1), 4),
            ]))
            story.append(info_table)
            story.append(Spacer(1, 16))

            # Items table
            tdata = [["#", "Item", "Qty", "Unit Price", "Tax", "Discount", "Total"]]
            for i, it in enumerate(items, 1):
                tdata.append([
                    str(i),
                    it.item_name or "",
                    str(it.quantity),
                    f"PKR {float(it.unit_price):,.2f}",
                    f"PKR {float(it.tax_amount):,.2f}",
                    f"PKR {float(it.discount_amount):,.2f}",
                    f"PKR {float(it.total):,.2f}"
                ])

            col_widths = [0.3*inch, 2.2*inch, 0.5*inch, 1.2*inch, 1*inch, 1*inch, 1.2*inch]
            items_table = Table(tdata, colWidths=col_widths, repeatRows=1)
            items_table.setStyle(TableStyle([
                ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#1e3a5f")),
                ("TEXTCOLOR", (0,0), (-1,0), colors.white),
                ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
                ("FONTSIZE", (0,0), (-1,-1), 9),
                ("ALIGN", (2,0), (-1,-1), "RIGHT"),
                ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#f8fafc")]),
                ("GRID", (0,0), (-1,-1), 0.5, colors.HexColor("#e2e8f0")),
                ("TOPPADDING", (0,0), (-1,-1), 6),
                ("BOTTOMPADDING", (0,0), (-1,-1), 6),
            ]))
            story.append(items_table)
            story.append(Spacer(1, 16))

            # Totals
            totals_data = [
                ["", "Subtotal:", f"PKR {float(inv.subtotal or 0):,.2f}"],
                ["", "Tax:", f"PKR {float(inv.tax_amount or 0):,.2f}"],
                ["", "Discount:", f"- PKR {float(inv.discount_amount or 0):,.2f}"],
                ["", "GRAND TOTAL:", f"PKR {float(inv.total_amount or 0):,.2f}"],
            ]
            totals_table = Table(totals_data, colWidths=[3.5*inch, 1.5*inch, 1.5*inch])
            totals_table.setStyle(TableStyle([
                ("ALIGN", (1,0), (-1,-1), "RIGHT"),
                ("FONTNAME", (1,3), (-1,3), "Helvetica-Bold"),
                ("FONTSIZE", (0,0), (-1,-1), 10),
                ("FONTSIZE", (1,3), (-1,3), 12),
                ("TEXTCOLOR", (1,3), (-1,3), colors.HexColor("#1e3a5f")),
                ("TOPPADDING", (0,3), (-1,3), 8),
                ("LINEABOVE", (1,3), (-1,3), 1.5, colors.HexColor("#1e3a5f")),
            ]))
            story.append(totals_table)

            story.append(Spacer(1, 24))
            footer_style = ParagraphStyle("footer", parent=styles["Normal"],
                                          fontSize=8, textColor=colors.HexColor("#94a3b8"),
                                          alignment=1)
            story.append(Paragraph("Thank you for your business.", footer_style))

            doc.build(story)
            return pdf_path

        except ImportError:
            # Fallback: plain text "PDF"
            with open(pdf_path, "w") as f:
                f.write(f"INVOICE: {inv.invoice_number}\n")
                f.write(f"Lead: {lead.lead_name if lead else 'N/A'}\n")
                f.write(f"Total: PKR {float(inv.total_amount or 0):,.2f}\n")
                f.write(f"Status: {inv.status}\n")
                f.write("\nItems:\n")
                for it in items:
                    f.write(f"  - {it.item_name}: PKR {float(it.total):,.2f}\n")
            return pdf_path

    # ─────────────── CONTRACT SERVICE ───────────────
    @staticmethod
    def get_contracts(db: Session, lead_id: Optional[int] = None):
        q = db.query(Contract)
        if lead_id:
            q = q.filter(Contract.lead_id == lead_id)
        contracts = q.order_by(Contract.created_at.desc()).all()
        result = []
        for c in contracts:
            lead = db.query(Lead).filter(Lead.id == c.lead_id).first() if c.lead_id else None
            tmpl = db.query(ContractTemplate).filter(ContractTemplate.id == c.template_id).first() if c.template_id else None
            result.append({
                "id": c.id,
                "lead_id": c.lead_id,
                "lead_name": lead.lead_name if lead else None,
                "template_id": c.template_id,
                "template_name": tmpl.name if tmpl else None,
                "contract_number": c.contract_number,
                "file_path": c.file_path,
                "signed_file_path": c.signed_file_path,
                "status": c.status,
                "version": c.version,
                "created_at": c.created_at.isoformat() if c.created_at else None,
                "updated_at": c.updated_at.isoformat() if c.updated_at else None,
            })
        return result

    @staticmethod
    def get_contract(db: Session, contract_id: int):
        c = db.query(Contract).filter(Contract.id == contract_id).first()
        if not c: return None
        lead = db.query(Lead).filter(Lead.id == c.lead_id).first() if c.lead_id else None
        tmpl = db.query(ContractTemplate).filter(ContractTemplate.id == c.template_id).first() if c.template_id else None
        return {
            "id": c.id,
            "lead_id": c.lead_id,
            "lead_name": lead.lead_name if lead else None,
            "template_id": c.template_id,
            "template_name": tmpl.name if tmpl else None,
            "contract_number": c.contract_number,
            "file_path": c.file_path,
            "signed_file_path": c.signed_file_path,
            "status": c.status,
            "version": c.version,
            "created_at": c.created_at.isoformat() if c.created_at else None,
            "updated_at": c.updated_at.isoformat() if c.updated_at else None,
        }

    @staticmethod
    def update_contract(db: Session, contract_id: int, data: dict):
        c = db.query(Contract).filter(Contract.id == contract_id).first()
        if not c: return None
        for k, v in data.items():
            if hasattr(c, k):
                setattr(c, k, v)
        db.commit()
        db.refresh(c)
        return ERPService.get_contract(db, contract_id)

    @staticmethod
    def delete_contract(db: Session, contract_id: int):
        c = db.query(Contract).filter(Contract.id == contract_id).first()
        if not c: return False
        db.delete(c)
        db.commit()
        return True

    @staticmethod
    def generate_contract(db: Session, lead_id: int, template_id: Optional[int] = None, custom_content: Optional[str] = None):
        lead = db.query(Lead).filter(Lead.id == lead_id).first()
        template = db.query(ContractTemplate).filter(ContractTemplate.id == template_id).first() if template_id else None
        
        if not lead: return None
        if not template and not custom_content: return None

        lead_products = db.query(LeadProduct).filter(LeadProduct.lead_id == lead_id).all()
        total = sum(float(lp.total or 0) for lp in lead_products)

        products_text = "\n".join([
            f"- {lp.product_name} | Qty: {lp.quantity} | Price: PKR {float(lp.unit_price):,.2f} | Total: PKR {float(lp.total):,.2f}"
            for lp in lead_products
        ]) or "No products"

        placeholders = {
            "{{client_name}}": lead.lead_name or "",
            "{{company_name}}": lead.company_name or "",
            "{{email}}": lead.email_id or "",
            "{{phone}}": lead.mobile_no or lead.phone or "",
            "{{address}}": lead.address_line1 or "",
            "{{date}}": datetime.now().strftime("%d %B %Y"),
            "{{invoice_total}}": f"PKR {total:,.2f}",
            "{{products}}": products_text,
            "{{lead_id}}": str(lead.id),
            "{{lead_status}}": lead.status or "",
        }

        ts = datetime.now().strftime("%Y%m%d%H%M%S")
        contract_number = f"CON-{lead_id}-{ts}"
        os.makedirs("storage/contracts", exist_ok=True)

        # Check existing version count
        existing = db.query(Contract).filter(Contract.lead_id == lead_id).count()
        version = existing + 1

        file_path = f"storage/contracts/{contract_number}.docx"

        # Use python-docx if template is a real DOCX
        try:
            from docx import Document
            if template and template.file_path and os.path.exists(template.file_path):
                doc = Document(template.file_path)
                for para in doc.paragraphs:
                    for key, val in placeholders.items():
                        if key in para.text:
                            for run in para.runs:
                                run.text = run.text.replace(key, val)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                for key, val in placeholders.items():
                                    if key in para.text:
                                        for run in para.runs:
                                            run.text = run.text.replace(key, val)
                doc.save(file_path)
            else:
                # Create from scratch using custom_content or template content
                content = custom_content or (template.content if template else "Service Agreement Content")
                doc = Document()
                doc.add_heading(f"SERVICE AGREEMENT", 0)
                doc.add_heading(f"Contract: {contract_number}", level=1)
                doc.add_paragraph(f"Date: {placeholders['{{date}}']}")
                doc.add_paragraph(f"Client: {placeholders['{{client_name}}']}")
                doc.add_paragraph(f"Company: {placeholders['{{company_name}}']}")
                
                # Replace placeholders in custom content too
                for key, val in placeholders.items():
                    content = content.replace(key, val)
                
                doc.add_heading("Contract Terms", level=2)
                doc.add_paragraph(content)
                
                doc.add_heading("Products / Services", level=2)
                doc.add_paragraph(products_text)
                doc.add_heading("Financials", level=2)
                doc.add_paragraph(f"Total Contract Value: {placeholders['{{invoice_total}}']}")
                doc.save(file_path)
        except ImportError:

            # Fallback: save as text with .docx extension
            with open(file_path, "w") as f:
                f.write(f"CONTRACT: {contract_number}\n\n")
                for k, v in placeholders.items():
                    f.write(f"{k.strip('{}')}: {v}\n")

        db_contract = Contract(
            lead_id=lead_id,
            template_id=template_id,
            contract_number=contract_number,
            file_path=file_path,
            status="Draft",
            version=version
        )
        db.add(db_contract)
        lead.contract_status = "Generated"
        ERPService._log(db, lead_id=lead_id, action="Contract Generated",
                        entity_type="Contract", entity_id=0,
                        details=f"Contract {contract_number} v{version} generated")
        db.commit()
        db.refresh(db_contract)
        return db_contract

    # ─────────────── PAYMENT SERVICE ───────────────
    @staticmethod
    def create_payment(db: Session, payment_data: PaymentCreate):
        db_payment = Payment(**payment_data.model_dump())
        db.add(db_payment)
        invoice = db.query(Invoice).filter(Invoice.id == payment_data.invoice_id).first()
        if invoice:
            invoice.status = "Paid"
            if invoice.lead_id:
                lead = db.query(Lead).filter(Lead.id == invoice.lead_id).first()
                if lead:
                    lead.payment_status = "Paid"
                    ERPService._log(db, lead_id=lead.id, action="Payment Received",
                                    entity_type="Invoice", entity_id=invoice.id,
                                    details=f"Payment of PKR {payment_data.amount:,.2f} received via {payment_data.payment_method}")
        db.commit()
        db.refresh(db_payment)
        return db_payment

    # ─────────────── ACTIVITY SERVICE ───────────────
    @staticmethod
    def get_lead_activity_timeline(db: Session, lead_id: int):
        return db.query(Activity).filter(Activity.lead_id == lead_id).order_by(Activity.created_at.desc()).all()

    @staticmethod
    def _log(db: Session, lead_id: Optional[int], action: str,
             entity_type: str, entity_id: int, details: str, performed_by: Optional[str] = None):
        activity = Activity(
            lead_id=lead_id,
            performed_by=performed_by,
            action=action,
            entity_type=entity_type,
            entity_id=entity_id,
            details=details,
            created_at=datetime.now()
        )
        db.add(activity)

