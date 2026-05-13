from sqlalchemy.orm import Session
from app.models.contract import Contract
from app.models.contract_template import ContractTemplate
from app.models.lead import Lead
from app.models.activity import Activity
from docx import Document
from datetime import datetime
import os

STORAGE_PATH = "public/storage"

class ContractService:
    @staticmethod
    def generate_contract(db: Session, lead_id: int, template_id: int):
        lead = db.query(Lead).filter(Lead.id == lead_id).first()
        tmpl = db.query(ContractTemplate).filter(ContractTemplate.id == template_id).first()
        if not lead or not tmpl: return None, "Lead or Template not found"
        
        os.makedirs(f"{STORAGE_PATH}/contracts", exist_ok=True)
        fname = f"Contract_{lead_id}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        fpath = f"{STORAGE_PATH}/contracts/{fname}"
        
        # Real logic: Replace placeholders in tmpl.file_path
        # Simplified for now:
        doc = Document()
        doc.add_heading(tmpl.name, 0)
        doc.add_paragraph(f"Client: {lead.lead_name}")
        doc.add_paragraph(f"Company: {lead.company_name}")
        doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph("Terms and conditions placeholder...")
        doc.save(fpath)
        
        contract = Contract(
            lead_id=lead_id,
            template_id=template_id,
            status="Draft",
            generated_docx_path=fpath
        )
        db.add(contract)
        lead.contract_status = "Generated"
        db.add(Activity(lead_id=lead_id, action="Contract Generated", detail=f"Using template {tmpl.name}"))
        db.commit()
        return contract, None
